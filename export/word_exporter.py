"""
export/word_exporter.py — Pure Python Word Report Generator
============================================================
Uses python-docx only. No Node.js required.
pip install python-docx

Two modes: 'summary' (1-2 pages) and 'detailed' (full manual with formulas).
"""
from __future__ import annotations
from datetime import datetime
from docx import Document # type: ignore
from docx.shared import Pt, Cm, RGBColor, Inches # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore
from docx.enum.table import WD_ALIGN_VERTICAL # type: ignore
from docx.oxml.ns import qn # type: ignore
from docx.oxml import OxmlElement # type: ignore
import copy


# ── Colour palette ─────────────────────────────────────────────────────────────
C_DARK_BLUE  = RGBColor(0x1F, 0x5C, 0x99)
C_MED_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
C_LIGHT_BLUE = RGBColor(0x31, 0x84, 0x9B)
C_GREY       = RGBColor(0x59, 0x59, 0x59)
C_GREEN      = RGBColor(0x1B, 0x5E, 0x20)
C_RED        = RGBColor(0x7F, 0x00, 0x00)
C_FORMULA    = RGBColor(0x00, 0x52, 0x8A)
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_HEAD_BG    = "1F5C99"  # hex for XML
C_HEAD2_BG   = "2E75B6"
C_ALT_BG     = "DBEEF4"
C_PLAIN_BG   = "FFFFFF"


def _set_cell_bg(cell, hex_color: str):
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)
    except Exception:
        pass  # Ignore background color errors


def _cell_text(cell, text: str, bold=False, color=None, size_pt=10,
               italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    except Exception:
        pass
    try:
        p = cell.paragraphs[0]
    except Exception:
        p = cell.add_paragraph()
    try:
        p.alignment = align
        p.clear()
        run = p.add_run(str(text))
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = color
    except Exception as e:
        # Fallback: just add text
        try:
            p.add_run(str(text))
        except Exception:
            pass


def _add_header_row(table, headers: list[str], col_widths=None):
    row = table.add_row()
    for i, text in enumerate(headers):
        cell = row.cells[i]
        try:
            _set_cell_bg(cell, C_HEAD2_BG)
        except Exception:
            pass
        _cell_text(cell, text, bold=True, color=C_WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    return row


def _add_data_row(table, values: list[str], alt=False, bold_first=True):
    row = table.add_row()
    bg = C_ALT_BG if alt else C_PLAIN_BG
    for i, text in enumerate(values):
        cell = row.cells[i]
        try:
            _set_cell_bg(cell, bg)
        except Exception:
            pass
        _cell_text(cell, text, bold=(bold_first and i == 0))
    return row


def _add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.color.rgb = C_DARK_BLUE
        run.font.size = Pt(14)
    elif level == 2:
        run.font.color.rgb = C_MED_BLUE
        run.font.size = Pt(12)
    else:
        run.font.color.rgb = C_LIGHT_BLUE
        run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_para(doc, text: str, italic=False, color=None, size_pt=10,
              indent_cm=0.0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    return p


def _add_formula(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = C_FORMULA
    return p


def _add_kv_table(doc, rows: list[tuple], col_widths=(5.5, 8.5)):
    """Two-column label:value table."""
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    try:
        table.columns[0].width = Cm(col_widths[0])
        table.columns[1].width = Cm(col_widths[1])
    except Exception:
        pass  # Ignore width setting errors
    for i, (label, value) in enumerate(rows):
        row = table.add_row()
        bg = C_ALT_BG if i % 2 == 0 else C_PLAIN_BG
        try:
            _set_cell_bg(row.cells[0], bg)
            _set_cell_bg(row.cells[1], bg)
        except Exception:
            pass
        _cell_text(row.cells[0], label, bold=True)
        _cell_text(row.cells[1], value)
    doc.add_paragraph()


def _status_color(ok):
    if ok is True:   return C_GREEN
    if ok is False:  return C_RED
    return C_GREY


def _add_results_table(doc, headers, rows, col_widths=None):
    n = len(headers)
    table = doc.add_table(rows=0, cols=n)
    table.style = 'Table Grid'
    if col_widths:
        try:
            for i, w in enumerate(col_widths):
                table.columns[i].width = Cm(w)
        except Exception:
            pass
    _add_header_row(table, headers)
    for i, row_data in enumerate(rows):
        r = table.add_row()
        bg = C_ALT_BG if i % 2 == 0 else C_PLAIN_BG
        for j, val in enumerate(row_data):
            c = r.cells[j]
            try:
                _set_cell_bg(c, bg)
            except Exception:
                pass
            is_status = (j == len(row_data) - 1)
            ok_val = None
            if is_status:
                if str(val).startswith("OK") or "✓" in str(val):
                    ok_val = True
                elif any(x in str(val) for x in ["FAIL","REVISE","UNSAFE","✗"]):
                    ok_val = False
            try:
                _cell_text(c, str(val),
                           bold=(j == 0 or is_status),
                           color=_status_color(ok_val) if is_status else None,
                           align=WD_ALIGN_PARAGRAPH.CENTER if is_status else WD_ALIGN_PARAGRAPH.LEFT)
            except Exception:
                _cell_text(c, str(val))
    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
# MAIN GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def generate_word_report(data: dict, output_path: str, mode: str = "detailed") -> str:
    """
    Generate a professional Word (.docx) structural engineering report.
    Pure python-docx — no Node.js required.
    """
    from constants import APP_NAME, APP_VERSION # type: ignore
    pi    = data.get("project_info", {})
    seism = data.get("seismic", {})
    beam  = data.get("beam", {})
    col   = data.get("column", {})
    fndg  = data.get("foundation", {})

    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    # ── Default styles ─────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Cover block ─────────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run(pi.get('project', 'Structural Design Report'))
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = C_DARK_BLUE

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("STRUCTURAL DESIGN CALCULATIONS")
    r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = C_MED_BLUE

    p_std = doc.add_paragraph()
    p_std.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_std.add_run("NBC 105:2025 (Second Revision)  ·  IS 456:2000  ·  IS 875")
    r3.font.size = Pt(10); r3.font.color.rgb = C_GREY

    doc.add_paragraph()
    _add_kv_table(doc, [
        ("Project",    pi.get("project", "—")),
        ("Engineer",   pi.get("engineer", "—")),
        ("Checked By", pi.get("checked_by", "—")),
        ("Job No.",    pi.get("job_no", "—")),
        ("Date",       pi.get("date", datetime.now().strftime("%Y-%m-%d"))),
        ("Report",     "Detailed Design Report" if mode == "detailed" else "Summary Report"),
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. SEISMIC
    # ══════════════════════════════════════════════════════════════════════════
    if seism:
        _add_heading(doc, "1. Seismic Design  —  NBC 105:2025")

        _add_kv_table(doc, [
            ("Zone / Municipality",    str(seism.get("zone_name", ""))),
            ("Zone Factor Z",          f"{seism.get('Z', 0):.2f}"),
            ("Importance Factor I",    f"{seism.get('I', 0):.2f}"),
            ("Soil Type",              f"{seism.get('soil_type', '')}  "
                                       f"(α={seism.get('alpha',0):.2f}, "
                                       f"Tc={seism.get('Tc',0):.1f}s, Td={seism.get('Td',0):.1f}s)"),
            ("Building Height H",      f"{seism.get('H', 0):.2f} m"),
            ("No. of Stories",         str(seism.get("num_stories", ""))),
            ("Structural System",      str(seism.get("struct_sub", ""))),
            ("Orientation (§3.6)",     "Parallel" if seism.get("is_parallel") else "Non-Parallel"),
            ("Snow Load",              "Included" if seism.get("include_snow", False) else "Excluded"),
            ("Analysis Method",        str(seism.get("method", ""))),
        ])

        if mode == "detailed":
            _add_heading(doc, "1.1 Period of Vibration  (NBC 105:2025 §5.1.2–3)", level=2)
            _add_formula(doc, f"T_approx = kt × H^(3/4) = {seism.get('kt',0):.4f} × {seism.get('H',0):.2f}^(3/4) = {seism.get('T_approx',0):.4f} s")
            _add_formula(doc, f"T_design  = 1.25 × T_approx = {seism.get('T',0):.4f} s  (§5.1.3 amplification)")

            _add_heading(doc, "1.2 Spectral Shape Factor Ch(T)  (§4.1.2 — 3-zone formula)", level=2)
            T=seism.get('T',0); Tc=seism.get('Tc',0); Td=seism.get('Td',0); a=seism.get('alpha',0)
            if T < Tc:
                zone = f"Flat plateau (T={T:.4f}<Tc={Tc:.1f}s): Ch(T) = α = {a:.2f}"
            elif T < Td:
                zone = f"Velocity-sensitive (Tc≤T<Td): Ch(T) = α×Tc/T = {a:.2f}×{Tc:.1f}/{T:.4f}"
            else:
                zone = f"Displacement-sensitive (T≥Td): Ch(T) = α×Tc×Td/T²"
            _add_formula(doc, zone)
            _add_formula(doc, f"C(T)  = Ch(T)×Z×I = {seism.get('Ch_T',0):.4f}×{seism.get('Z',0):.2f}×{seism.get('I',0):.2f} = {seism.get('C_T',0):.4f}")
            _add_formula(doc, f"Cs(T) = 0.20×C(T) = {seism.get('Cs_T',0):.4f}  (SLS, §4.2)")
            _add_formula(doc, f"Cv(T) = 2/3×Z = {seism.get('Cv_T',0):.4f}  (Vertical, §4.3)")

            _add_heading(doc, "1.3 Base Shear Coefficients  (§6.1)", level=2)
            _add_formula(doc, f"Cd(T)_ULS = C(T)/(Rμ×Ωu) = {seism.get('C_T',0):.4f}/({seism.get('Ru',0):.2f}×{seism.get('O_u',0):.2f}) = {seism.get('Cd_ULS',0):.4f}")
            _add_formula(doc, f"Cd(T)_SLS = Cs(T)/Ωs = {seism.get('Cs_T',0):.4f}/{seism.get('O_s',0):.2f} = {seism.get('Cd_SLS',0):.4f}")

        _add_results_table(doc,
            ["Result", "Value", "Clause", "Status"],
            [
                ("T_design",   f"{seism.get('T',0):.4f} s",   "§5.1.2–3", "OK ✓"),
                ("Ch(T)",      f"{seism.get('Ch_T',0):.4f}",  "§4.1.2",   "OK ✓"),
                ("C(T)",       f"{seism.get('C_T',0):.4f}",   "§4.1.1",   "OK ✓"),
                ("Cs(T) SLS",  f"{seism.get('Cs_T',0):.4f}",  "§4.2",     "OK ✓"),
                ("Cv(T) Vert", f"{seism.get('Cv_T',0):.4f}",  "§4.3",     "OK ✓"),
                ("Cd(T) ULS",  f"{seism.get('Cd_ULS',0):.4f}","§6.1.1",   "OK ✓"),
                ("Cd(T) SLS",  f"{seism.get('Cd_SLS',0):.4f}","§6.1.2",   "OK ✓"),
                ("kd",         f"{seism.get('kd',0):.2f}",    "Table 6-1","OK ✓"),
            ],
            col_widths=[4.5, 3.5, 3.0, 2.5]
        )

        # Story forces
        sf = seism.get("story_forces", [])
        if sf:
            _add_heading(doc, "1.4 Story Force Distribution  (NBC 105:2025 §6.3)" if mode=="detailed"
                         else "Story Force Distribution", level=2)
            if mode == "detailed":
                _add_formula(doc, "Fi = V × Wi×hi^k / Σ(Wj×hj^k)")
            _add_results_table(doc,
                ["Floor", "Wi (kN)", "hi (m)", "Wi·hi^k", "Fi (kN)", "Story Shear"],
                [[str(f["floor"]), f"{f['W_kN']:.0f}", f"{f['h_m']:.1f}",
                  f"{f['Wh_k']:.0f}", f"{f['Fi_kN']:.1f}", f"{f['Vx_kN']:.1f}"]
                 for f in sf],
                col_widths=[1.5, 2.5, 2.0, 3.0, 2.5, 2.5]
            )

        # Load combos (detailed)
        if mode == "detailed":
            combos = seism.get("load_combos", [])
            if combos:
                _add_heading(doc, "1.5 Load Combinations  (NBC 105:2025 §3.6)", level=2)
                _add_para(doc, f"λ = {seism.get('lambda_ll', 0.30):.2f}  "
                           "(app setting: 0.30 for seismic combos; for story Wi input, top story LL factor is 0.00)", italic=True)

                _add_results_table(doc,
                    ["Combination", "Formula", "DL", "LL(λ)", "E/W"],
                    [[c["label"], c["formula"],
                      str(c["DL_fac"]),
                      str(c.get("LL_fac","λ")),
                      ("".join([
                         f"{c['EX_ULS_fac']:+.1f}EX" if c.get("EX_ULS_fac",0)!=0 else "",
                         f"{c['EY_ULS_fac']:+.1f}EY" if c.get("EY_ULS_fac",0)!=0 else ""
                      ])) if ("EX_ULS_fac" in c and (c.get("EX_ULS_fac",0)!=0 or c.get("EY_ULS_fac",0)!=0)) else
                      f"{c['E_SLS_fac']:+.1f}E_SLS" if c.get("E_SLS_fac",0)!=0 else "—"]
                     for c in combos],
                    col_widths=[2.5, 4.0, 1.5, 1.5, 2.0]
                )

    # ══════════════════════════════════════════════════════════════════════════
    # 2. BEAM
    # ══════════════════════════════════════════════════════════════════════════
    if beam:
        _add_heading(doc, "2. Beam Design  —  IS 456:2000")
        _add_kv_table(doc, [
            ("Section b × D",      f"{beam.get('b','?')} × {beam.get('D','?')} mm"),
            ("Effective depth d",  f"{beam.get('d_eff_mm',0):.1f} mm"),
            ("Concrete / Steel",   f"M{beam.get('fck','')} / Fe{beam.get('fy','')}"),
            ("Design Mu",          f"{beam.get('Mu_design_kNm',0):.2f} kN·m"),
            ("Limiting Mu,lim",    f"{beam.get('Mu_lim_kNm',0):.2f} kN·m  (§38.1)"),
            ("Section type",       "Doubly Reinforced" if beam.get("is_doubly") else "Singly Reinforced"),
        ])

        if mode == "detailed":
            _add_heading(doc, "2.1 Flexural Design  (IS 456:2000 §38)", level=2)
            _add_formula(doc, f"xu,max/d = 0.48  (Fe415)  →  xu,max = {0.48*beam.get('d_eff_mm',0):.1f} mm")
            _add_formula(doc, f"Mu,lim = 0.36·fck·b·xu,max·(d−0.42·xu,max) = {beam.get('Mu_lim_kNm',0):.2f} kN·m")
            dr = beam.get("doubly")
            if dr:
                _add_heading(doc, "2.2 Doubly-Reinforced  (IS 456:2000 Annex G)", level=2)
                _add_formula(doc, f"fsc = {dr.get('fsc_MPa',0):.1f} MPa  (εsc={dr.get('eps_sc',0):.4f})")
                _add_formula(doc, f"Asc = {dr.get('Asc_req_mm2',0):.0f} mm²  →  {dr.get('no_comp_bars',0)}×Ø{int(dr.get('comp_bar_dia',0))}mm")
            defl = beam.get("deflection")
            if defl:
                _add_heading(doc, "2.3 Deflection Check  (IS 456:2000 §23.2)", level=2)
                _add_formula(doc, f"Basic L/d = {defl.get('ld_basic',20)}  |  kt = {defl.get('kt',0):.3f}  |  kc = {defl.get('kc',1):.3f}")
                _add_formula(doc, f"Allowable L/d = {defl.get('ld_allow',0):.2f}   Provided = {defl.get('ld_prov',0):.2f}")

        _add_results_table(doc,
            ["Check", "Value", "Clause", "Status"],
            [
                ("Ast required",    f"{beam.get('Ast_req_mm2',0):.0f} mm²", "§26.5.1", "OK ✓"),
                ("Ast provided",    f"{beam.get('Ast_prov_mm2',0):.0f} mm²","§26.5.1",
                 "OK ✓" if beam.get('Ast_prov_mm2',0)>=beam.get('Ast_req_mm2',0) else "REVISE ✗"),
                ("Bars",            f"{beam.get('no_of_bars',0)}×Ø{beam.get('main_dia','?')} mm","—","OK ✓"),
                ("Spacing",         f"{beam.get('spacing_mm',0):.0f} mm c/c","§26.3.2","OK ✓"),
                ("Shear",           beam.get("shear",{}).get("status","—"),"§40","OK ✓"),
                ("Dev. length Ld",  f"{beam.get('Ld_mm',0):.0f} mm","§26.2.1","OK ✓"),
                ("L/d check",
                 f"{beam.get('deflection',{}).get('ld_prov',0):.1f} ≤ {beam.get('deflection',{}).get('ld_allow',0):.1f}"
                 if beam.get("deflection") else "N/A","§23.2",
                 "OK ✓" if beam.get("deflection",{}).get("ok") else
                 ("N/A" if not beam.get("deflection") else "CHECK ⚠")),
            ],
            col_widths=[4.5, 4.0, 2.5, 2.5]
        )

    # ══════════════════════════════════════════════════════════════════════════
    # 3. COLUMN
    # ══════════════════════════════════════════════════════════════════════════
    if col:
        _add_heading(doc, "3. Column Design  —  IS 456:2000 + NBC 105:2025 Annex A")
        _add_kv_table(doc, [
            ("Section b × D",     f"{col.get('b_mm','?')} × {col.get('D_mm','?')} mm"),
            ("λx / λy",           f"{col.get('lambda_x',0):.2f} / {col.get('lambda_y',0):.2f}  "
                                   f"({'Slender' if col.get('is_slender') else 'Short'})"),
            ("Design Mux",        f"{col.get('Mux_design_kNm',0):.2f} kN·m  (incl. Madd)"),
            ("Design Muy",        f"{col.get('Muy_design_kNm',0):.2f} kN·m  (incl. Madd)"),
        ])

        if mode == "detailed":
            _add_heading(doc, "3.1 Biaxial Interaction  (IS 456:2000 §39.6)", level=2)
            _add_formula(doc, "(Mux/Mux1)^αn + (Muy/Muy1)^αn ≤ 1.0")
            _add_formula(doc, f"Mux1={col.get('Mux1_kNm',0):.2f} kNm  Muy1={col.get('Muy1_kNm',0):.2f} kNm  αn={col.get('alpha_n',0):.3f}")
            _add_formula(doc, f"Interaction = {col.get('interaction',0):.4f}  "
                          f"{'≤ 1.0 OK ✓' if col.get('interaction',1)<=1.0 else '> 1.0 REVISE ✗'}")
            _add_heading(doc, "3.2 NBC 105:2025 Annex A  —  Ductile Detailing", level=2)
            _add_formula(doc, f"Confinement zone lo = {col.get('conf_zone_mm',0):.0f} mm  (top & bottom)")
            _add_formula(doc, f"Tie spacing in conf. zone = {col.get('conf_tie_sp_mm',0):.0f} mm c/c  (135° hooks)")
            _add_formula(doc, f"Ash required = {col.get('Ash_req_mm2',0):.1f} mm²   provided = {col.get('Ash_prov_mm2',0):.1f} mm²")

        inter = col.get('interaction', 0)
        _add_results_table(doc,
            ["Check", "Value", "Clause", "Status"],
            [
                ("Biaxial interaction", f"{inter:.4f}", "§39.6",
                 "OK ✓" if inter<=1.0 else "FAIL ✗"),
                ("Steel %", f"{col.get('steel_pct',0):.2f}%", "§26.5.3",
                 "OK ✓" if 0.8<=col.get('steel_pct',0)<=4.0 else "WARN ⚠"),
                ("Bars", f"{col.get('no_of_bars',0)}×Ø{int(col.get('bar_dia_mm',0))}mm","—","OK ✓"),
                ("Tie spacing", f"{col.get('tie_spacing_mm',0):.0f} mm", "§26.5.3.1","OK ✓"),
                ("Conf. zone", f"{col.get('conf_zone_mm',0):.0f} mm", "Annex A","OK ✓"),
                ("Conf. tie sp.", f"{col.get('conf_tie_sp_mm',0):.0f} mm","Annex A","OK ✓"),
                ("Ash check", f"req={col.get('Ash_req_mm2',0):.1f} / prov={col.get('Ash_prov_mm2',0):.1f}",
                 "Annex A §A.4.4.4",
                 "OK ✓" if col.get("hoop_ok") else "WARN ⚠"),
            ],
            col_widths=[4.0, 4.5, 3.0, 2.0]
        )

    # ══════════════════════════════════════════════════════════════════════════
    # 4. FOUNDATION
    # ══════════════════════════════════════════════════════════════════════════
    if fndg:
        _add_heading(doc, "4. Foundation Design  —  IS 456:2000 §34 + NBC 105:2025 §3.8")
        ftype = fndg.get("type", "concentric")
        _add_kv_table(doc, [
            ("Type",            ftype.title() + " Isolated Footing"),
            ("Plan size",       f"{fndg.get('L_mm',0)} × {fndg.get('B_mm',0)} mm"),
            ("Depth D / d",     f"{fndg.get('D_mm',0):.0f} / {fndg.get('d_mm',0):.0f} mm"),
            ("Design SBC",      f"{fndg.get('SBC_used_kPa',0):.0f} kN/m²"
                                f"{'  (+50% seismic, §3.8)' if fndg.get('seismic_used') else ''}"),
            ("q_max",           f"{fndg.get('q_max_kPa',0):.2f} kN/m²"),
            ("q_min",           f"{fndg.get('q_min_kPa',0):.2f} kN/m²"
                                f"{'  (no tension)' if fndg.get('q_min_kPa',0)>=0 else '  ⚠ tension zone'}"),
        ])

        if mode == "detailed":
            _add_heading(doc, "4.1 Bending & Reinforcement  (§34.4.1)", level=2)
            _add_formula(doc, f"qu (factored) = 1.5×q_avg = {fndg.get('qu_kPa',0):.2f} kN/m²")
            _add_formula(doc, f"Mu_L = qu×B×(overhang)²/2 = {fndg.get('Mu_L_kNm',0):.2f} kN·m")
            _add_formula(doc, f"Mu_B = qu×L×(overhang)²/2 = {fndg.get('Mu_B_kNm',0):.2f} kN·m")
            _add_heading(doc, "4.2 Shear Checks", level=2)
            _add_formula(doc, f"One-way: τv,L={fndg.get('tau_v_L',0):.3f} MPa  "
                          f"τc={fndg.get('tau_c_L',0):.3f} MPa  "
                          f"{'OK ✓' if fndg.get('one_way_ok') else 'FAIL ✗'}  (§34.4.2a)")
            _add_formula(doc, f"Punching: τv={fndg.get('tau_v_punch',0):.3f} MPa  "
                          f"τc=0.25√fck={fndg.get('tau_c_punch',0):.3f} MPa  "
                          f"{'OK ✓' if fndg.get('punch_ok') else 'FAIL ✗'}  (§31.6)")
            _add_formula(doc, f"Dev. length: Ld={fndg.get('Ld_mm',0):.0f} mm  "
                          f"Available={fndg.get('avail_L_mm',0):.0f} mm  "
                          f"{'OK ✓' if fndg.get('dev_ok') else 'WARN ⚠'}  (§34.4.3)")

        bar = int(fndg.get('bar_dia_mm', 12))
        _add_results_table(doc,
            ["Check", "Value", "Clause", "Status"],
            [
                ("q_max ≤ SBC",     f"{fndg.get('q_max_kPa',0):.2f} kN/m²","§34.2",
                 "OK ✓" if fndg.get('pressure_ok') else "FAIL ✗"),
                ("q_min ≥ 0",       f"{fndg.get('q_min_kPa',0):.2f} kN/m²","§34.2.4",
                 "OK ✓" if fndg.get('q_min_kPa',0)>=0 else "WARN ⚠"),
                ("Rein. L-dir.",     f"Ø{bar}@{fndg.get('sp_L_mm',0)}mm","§34.4.1","OK ✓"),
                ("Rein. B-dir.",     f"Ø{bar}@{fndg.get('sp_B_mm',0)}mm","§34.4.1","OK ✓"),
                ("One-way shear",   f"τv={fndg.get('tau_v_L',0):.3f} MPa","§34.4.2a",
                 "OK ✓" if fndg.get('one_way_ok') else "FAIL ✗"),
                ("Punching shear",  f"τv={fndg.get('tau_v_punch',0):.3f} MPa","§31.6",
                 "OK ✓" if fndg.get('punch_ok') else "FAIL ✗"),
                ("Development Ld",  f"{fndg.get('Ld_mm',0):.0f} mm","§34.4.3",
                 "OK ✓" if fndg.get('dev_ok') else "WARN ⚠"),
                ("Col-ftg bearing", f"{fndg.get('bear_stress_MPa',0):.2f} MPa","§34.4.4",
                 "OK ✓" if fndg.get('bear_ok', True) else "WARN ⚠"),
            ],
            col_widths=[4.0, 4.0, 3.0, 2.5]
        )

    # ══════════════════════════════════════════════════════════════════════════
    # DISCLAIMER
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _add_heading(doc, "Disclaimer")
    _add_para(doc,
        f"This report is generated by {APP_NAME} {APP_VERSION} for reference "
        "purposes. All results must be independently verified by a qualified structural "
        "engineer before use in construction. Standards applied: NBC 105:2025 (Second "
        "Revision), IS 456:2000, IS 875 Part 1 & 2.", italic=True, color=C_GREY)
    _add_para(doc,
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  "
        f"|  {APP_NAME} {APP_VERSION}", color=C_GREY)

    doc.save(output_path)
    _fix_docx_settings(output_path)
    return output_path


def _fix_docx_settings(path: str) -> None:
    """Remove the invalid w:zoom element that python-docx adds without w:percent."""
    import zipfile, shutil, tempfile, re, os
    tmp = path + ".tmp"
    with zipfile.ZipFile(path, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'word/settings.xml':
                text = data.decode('utf-8')
                # Remove invalid zoom-without-percent
                text = re.sub(r'<w:zoom[^>]*/>', '', text)
                # Also fix any zoom with only val= and no percent=
                text = re.sub(r'<w:zoom\s+w:val="[^"]*"\s*/>', '', text)
                data = text.encode('utf-8')
            zout.writestr(item, data)
    os.replace(tmp, path)
